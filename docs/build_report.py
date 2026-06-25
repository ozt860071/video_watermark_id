"""Generate docs/report.docx — project report on the video watermarking benchmark.

Run:
    python docs/build_report.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


OUT = Path(__file__).resolve().parent / "report.docx"
IMG_DIR = Path(__file__).resolve().parent / "snapshots"

_PRESETS = [
    "recompress_crf28", "recompress_crf32", "recompress_crf36",
    "transcode_x265",
    "resize_0.75", "resize_0.5",
    "crop_0.8", "crop_0.6",
    "noise_sigma2", "noise_sigma5", "noise_sigma10",
    "frame_drop_0.5",
]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Helvetica"
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Helvetica"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Helvetica"
        run.font.size = Pt(size)


def add_table(doc, headers, rows, col_widths=None, font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Helvetica"
                run.font.size = Pt(font_size)
    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            cell = table.rows[r].cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Helvetica"
                    run.font.size = Pt(font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = w
    return table


def _add_snapshot_table(doc, img_dir: Path):
    """3-column grid: preset | DCT | TrustMark."""
    cell_img_w = Inches(2.4)
    baseline_w = Inches(4.9)

    table = doc.add_table(rows=1 + 1 + len(_PRESETS), cols=3)
    table.style = "Light Grid Accent 1"

    for i, txt in enumerate(["Preset", "DCT", "TrustMark"]):
        cell = table.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Helvetica"
                r.font.size = Pt(9)

    base_cells = table.rows[1].cells
    base_cells[0].text = "baseline\n(original)"
    for p in base_cells[0].paragraphs:
        for r in p.runs:
            r.font.name = "Helvetica"
            r.font.size = Pt(9)
    merged = base_cells[1].merge(base_cells[2])
    merged.paragraphs[0].add_run().add_picture(
        str(img_dir / "frame600_baseline.jpg"), width=baseline_w
    )

    for ri, preset in enumerate(_PRESETS, start=2):
        cells = table.rows[ri].cells
        cells[0].text = preset
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.font.name = "Helvetica"
                r.font.size = Pt(10)
        cells[1].paragraphs[0].add_run().add_picture(
            str(img_dir / f"frame600_dct_{preset}.jpg"), width=cell_img_w
        )
        cells[2].paragraphs[0].add_run().add_picture(
            str(img_dir / f"frame600_trustmark_{preset}.jpg"), width=cell_img_w
        )
    return table


def main():
    doc = Document()

    # Standard-ish margins (looser than the previous 0.7/0.8 version)
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    # ----- Title -----
    title = doc.add_heading("Video Watermarking Benchmark", level=0)
    for run in title.runs:
        run.font.name = "Helvetica"
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "DCT-QIM vs. Adobe TrustMark — quality, robustness, and failure-mode "
        "characterisation under codec, geometric, and screen-capture attacks."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ----- 1. Problem framing -----
    add_heading(doc, "1. Problem framing", level=1)
    add_para(
        doc,
        "We compare two video-watermarking approaches on a common task: "
        "embed a 56-bit identifier per frame, run a realistic codec "
        "round-trip and a downstream attack suite, then measure how often "
        "the payload can be recovered. The two methods are a custom "
        "DCT-QIM (classical frequency-domain) scheme and Adobe TrustMark "
        "(learned encoder/decoder pair). The benchmark answers:"
    )
    add_bullets(doc, [
        "How invisible is each watermark on real content? (PSNR, SSIM)",
        "How robust is each after H.264 / H.265 plus attacks? (per-frame BER, exact-match rate)",
        "Which frames are easier to recover from, and why?",
    ])

    # ----- 2. Assumptions and constraints -----
    add_heading(doc, "2. Assumptions and constraints", level=1)
    add_bullets(doc, [
        "Payload is 56 bits (7 bytes). Matches DCT's BCH(127, t=11) data "
        "field and TrustMark's BCH_5 61-bit capacity.",
        "Per-frame, independent embedding. The decoder is free to "
        "majority-vote across frames after the fact.",
        "Encoder and decoder share a secret key and assume the original "
        "frame geometry — pixel-level scale/translate breaks the DCT path.",
        "Reference codec: H.264 (libx264) at CRF 23; H.265 (libx265) used "
        "for cross-codec stress.",
        "Test clip: 1920×800 cinematic excerpt from Tears of Steel, "
        "1432 frames (≈60 s, 24 fps). Black-intro frames excluded.",
    ])

    # ----- 3. Methods (encode + decode) -----
    add_heading(doc, "3. Methods", level=1)

    add_heading(doc, "3.1 DCT-QIM", level=2)
    add_para(
        doc,
        "The frame is converted to BT.601 full-range YCbCr; embedding "
        "happens only in the Y (luma) plane, since H.264/265 chroma is "
        "aggressively subsampled and quantised. The Y plane is tiled into "
        "8×8 blocks (the same grid the codec uses for transform coding) "
        "and a 2-D orthonormal DCT-II is applied per block."
    )
    add_para(
        doc,
        "The 56-bit payload is expanded by BCH(n=127, k=57, t=11) into a "
        "127-chip codeword. Each chip is embedded in a pseudo-randomly "
        "chosen block via Quantisation Index Modulation (QIM) of a "
        "mid-band AC coefficient (zigzag indices 5–14, where coefficients "
        "are large enough to survive moderate quantisation but small "
        "enough that perturbations remain visually imperceptible). The "
        "block-selection sequence is keyed by SHA-256 of the user secret. "
        "With num_blocks=4 each chip lands in four distinct blocks, "
        "providing an inner majority-vote on decode in addition to the "
        "outer BCH correction."
    )
    add_para(
        doc,
        "Decode mirrors encode: re-tile, DCT, read each QIM coefficient, "
        "soft-mean across blocks → 127 hard chips → BCH-decode → 56-bit "
        "payload. The decoder also returns BCH n_errors_corrected (or −1 "
        "for an uncorrectable codeword), useful as a per-frame confidence "
        "signal. QIM step δ = 30 is calibrated empirically to survive "
        "H.264 down to CRF ≈ 28 at PSNR ≈ 57 dB on natural content."
    )

    add_heading(doc, "3.2 TrustMark", level=2)
    add_para(
        doc,
        "TrustMark (Adobe, 2023) is a learned encoder/decoder pair. The "
        "encoder operates internally at 256×256 and re-projects a learned "
        "residual onto the source frame; the decoder is trained to be "
        "robust to scale, crop, noise, and codec compression within the "
        "digital channel. We use BCH_5 mode (61-bit capacity, corrects up "
        "to 5 bit-flips out of 100) — the most-robust mode that fits a "
        "56-bit payload."
    )
    add_para(
        doc,
        "The 56-bit payload is serialised as a 56-character binary string "
        "and passed to TrustMark.encode() with MODE='binary'. The "
        "per-frame adapter wraps the model with the same encode/decode "
        "signature as the DCT scheme so the benchmark harness is shared. "
        "Decode returns the recovered bits plus a watermark-present flag "
        "(BCH decode succeeded ≠ payload matched expected — see §8)."
    )

    # ----- 4. Attack model -----
    add_heading(doc, "4. Attack model", level=1)
    add_para(
        doc,
        "Attacks are applied after the main encode + codec round-trip. Six "
        "families × twelve presets cover the distortions seen in real "
        "distribution pipelines (re-streaming, mobile playback, format "
        "conversion). The attack model is non-adversarial — no "
        "watermark-aware adversary, no collusion or oracle attacks."
    )
    add_table(doc,
        headers=["Family", "Presets", "Stresses"],
        rows=[
            ["Re-encode",     "recompress_crf28 / 32 / 36",     "Codec quantisation severity"],
            ["Transcode",     "transcode_x265 (CRF 28)",        "Different codec / Q-matrix"],
            ["Spatial resize","resize_0.75 / 0.5",              "Down→up rescale"],
            ["Spatial crop",  "crop_0.8 / 0.6",                 "Center-crop → rescale to original"],
            ["Gaussian noise","noise_sigma2 / 5 / 10",          "Sensor noise, generation loss"],
            ["Frame drop",    "frame_drop_0.5",                 "Keep every other frame"],
        ],
        col_widths=[Inches(1.4), Inches(2.4), Inches(3.0)],
    )

    # ----- 5. Evaluation methodology -----
    add_heading(doc, "5. Evaluation methodology", level=1)
    add_para(
        doc,
        "One CLI invocation drives the whole pipeline "
        "(python -m video_watermark.benchmark.run). Per frame:"
    )
    add_bullets(doc, [
        "Load the original frame as uint8 RGB via PyAV.",
        "Run the chosen watermarker's encode(frame, payload_bits); "
        "watermarked frames stay in memory.",
        "Write all watermarked frames as an MP4 at the chosen codec/CRF "
        "— this is the main codec round-trip.",
        "Re-read the watermarked MP4, decode each frame, and compute "
        "per-frame quality (PSNR / SSIM / Δpx) and robustness "
        "(BER, exact-match, n_errors_corrected) metrics.",
        "If --attacks is given, for each preset: apply the attack to the "
        "watermarked MP4, decode, and emit a per-frame trace "
        "(frame_idx, ber, success) alongside aggregate statistics.",
    ])
    add_para(
        doc,
        "All numbers — including the full per-frame attack trace — land "
        "in benchmark_results.json for downstream analysis. The per-frame "
        "trace is what makes the 'which frames are easier?' question "
        "answerable rather than averaged-away."
    )

    # ----- 6. Metrics -----
    add_heading(doc, "6. Metrics", level=1)
    add_para(doc, "Quality (watermarked vs. original frame):", bold=True)
    add_bullets(doc, [
        "PSNR (dB) — target ≥ 42 dB for invisibility.",
        "SSIM — channel-averaged; target ≥ 0.98.",
        "Mean |Δpx| — average per-pixel absolute difference (0–255 scale).",
    ])
    add_para(doc, "Robustness (decoded vs. expected payload):", bold=True)
    add_bullets(doc, [
        "Per-frame BER — fraction of 56 bits that differ (0 = perfect, "
        "0.5 = destroyed).",
        "Exact-match rate (recovered / total) — headline number in tables.",
        "Per-bit BER — error rate per individual payload bit across frames.",
        "Majority-vote BER — soft mean of decoded bits across frames; "
        "models a video-level decoder that aggregates frames.",
        "BCH n_errors_corrected (DCT) — chip errors fixed per frame "
        "(or −1 = uncorrectable).",
    ])

    # ----- 7. Experimental results -----
    add_heading(doc, "7. Experimental results", level=1)
    add_para(
        doc,
        "Full-duration run, 1432 frames (≈60 s, 24 fps), 1920×800 source, "
        "payload \"ABCDEFG\" (0x41424344454647), main pipeline at H.264 "
        "CRF 23. Quality is similar between methods; TrustMark trades "
        "~10 dB of PSNR for substantially better robustness:"
    )
    add_table(doc,
        headers=["Metric", "DCT", "TrustMark"],
        rows=[
            ["PSNR mean / min (dB)",     "56.86 / 56.00",  "47.20 / 42.39"],
            ["SSIM mean",                "0.9986",         "0.9965"],
            ["Mean |Δpx|",               "0.033",          "0.488"],
            ["Baseline exact-match",     "96.7 %",         "90.85 %"],
            ["Baseline mean BER",        "0.0051",         "0.0310"],
        ],
        col_widths=[Inches(2.2), Inches(1.7), Inches(1.7)],
    )
    add_para(
        doc,
        "Attack-suite results — frames perfectly recovered out of 1432 "
        "(or 716 for frame_drop_0.5, which keeps every other frame):"
    )
    add_table(doc,
        headers=["Attack preset", "DCT ok/N", "DCT BER", "TM ok/N", "TM BER"],
        rows=[
            ["recompress_crf28", "604/1432",  "0.118", "1267/1432", "0.039"],
            ["recompress_crf32", "147/1432",  "0.211", "1068/1432", "0.086"],
            ["recompress_crf36", "1/1432",    "0.311", "626/1432",  "0.190"],
            ["transcode_x265",   "238/1432",  "0.199", "1193/1432", "0.057"],
            ["resize_0.75",      "489/1432",  "0.141", "1301/1432", "0.031"],
            ["resize_0.5",       "0/1432",    "0.340", "1300/1432", "0.031"],
            ["crop_0.8",         "0/1432",    "0.342", "1262/1432", "0.040"],
            ["crop_0.6",         "0/1432",    "0.339", "0/1432",    "0.341"],
            ["noise_sigma2",     "1214/1432", "0.037", "1298/1432", "0.032"],
            ["noise_sigma5",     "1151/1432", "0.043", "1299/1432", "0.032"],
            ["noise_sigma10",    "1070/1432", "0.050", "1291/1432", "0.034"],
            ["frame_drop_0.5",   "615/716",   "0.015", "650/716",   "0.031"],
        ],
        col_widths=[Inches(1.6), Inches(1.1), Inches(0.9),
                    Inches(1.1), Inches(0.9)],
    )
    add_bullets(doc, [
        "TrustMark dominates every codec and geometric attack except "
        "crop_0.6 (total failure for both) and Gaussian noise (where DCT "
        "is competitive).",
        "DCT collapses at resize_0.5 / crop_0.8 / crop_0.6 (all 0/1432) — "
        "its block grid loses alignment with the embedded positions.",
        "DCT recompress has a clean monotonic curve: 42 % → 10 % → ~0 % "
        "as CRF rises 28 → 32 → 36. The 56 dB PSNR headroom buys "
        "robustness up to about CRF 30, after which the codec quantiser "
        "starts erasing mid-band coefficients faster than BCH can recover.",
        "frame_drop_0.5 is the easiest attack for both: surviving frames "
        "are essentially the baseline, since frame drop is not a "
        "per-frame distortion.",
        "Per-frame observation enabled by the trace: under "
        "recompress_crf28, transcode_x265 and resize_0.75, DCT recovers "
        "isolated frames inside otherwise-failing runs — these tend to "
        "be I-frames, which the codec quantises more leniently. "
        "Aggregate BER alone would have hidden this I-frame vs P-frame "
        "robustness gap.",
    ])
    add_para(
        doc,
        "Compute footprint: end-to-end wall time was ~30 minutes on Apple "
        "Silicon (CPU TrustMark inference dominated; MPS support was "
        "added after this run — see §10). DCT encode ran at ~31 ms/frame, "
        "TrustMark encode at ~41 ms/frame; per-attack decode took 10–60 s "
        "(DCT) or 60–90 s (TrustMark) for the full 1432 frames. The "
        "attack-MP4 outputs totalled ~3.3 GB, dominated by the "
        "noise-attack files (~440 MB each — Gaussian noise defeats "
        "H.264's predictive compression).",
        italic=True,
    )

    # ----- 7.1 Camera capture test -----
    add_heading(doc, "7.1 Camera capture test (out-of-channel)", level=2)
    add_para(
        doc,
        "As a real-world stress test we played the DCT/TrustMark "
        "watermarked output on a laptop screen and re-captured it with an "
        "iPhone 15 Pro — two recordings, 1920×1080 at 30 fps (1816 frames) "
        "and 60 fps (3586 frames). Both decoders score 0 frames recovered "
        "out of every frame in both clips; per-frame BER sits pinned at the "
        "BCH-uncorrectable / no-signal floor (0.339). TrustMark's "
        "watermark-present flag triggers on ~1–2 % of frames, but the "
        "payload is garbage in every one of them (consistent with the "
        "BCH-5 false-positive rate of ≈0.2 % plus low-grade noise). The "
        "display→camera optical channel layers refresh-rate beating (laptop "
        "60 Hz vs camera 30 / 60 fps), moire from screen-pixel × Bayer "
        "interaction, rolling-shutter skew, glare, colour-profile drift "
        "and a fresh H.264 encode — collectively the print-scan / "
        "physical-channel attack, which neither model was trained for."
    )
    add_para(doc, "Possible improvements to close this gap:", bold=True)
    add_bullets(doc, [
        "Use a TrustMark variant trained with screen-capture / print-scan "
        "augmentations in its degradation set (Adobe ships such variants "
        "outside the default BCH_5 model used here).",
        "Add a synchronisation marker (e.g. a known-pattern border or "
        "corner glyph) so the decoder can geometrically rectify and "
        "deshake before passing the frame to the watermark decoder.",
        "Accumulate evidence across frames: even at ~2 % per-frame "
        "signal-trigger, a few seconds of capture provides enough soft "
        "votes for a structured majority decoder to recover the payload, "
        "provided the per-frame errors are independent.",
        "Drop frame rate to capture-aligned values (24 fps source → 24 / "
        "48 / 60 fps capture, with shutter speed tuned to the screen "
        "refresh) to reduce refresh beating and rolling-shutter skew.",
    ])

    # ----- 7.2 Visual reference -----
    add_heading(doc, "7.2 Visual reference (frame 600, ≈25 s)", level=2)
    add_para(
        doc,
        "Frame 600 (~25 s into the clip) from each attacked output, both "
        "methods side by side. Baseline (top) is the un-watermarked "
        "original; for frame_drop_0.5 the equivalent kept frame (output "
        "index 300) is shown so the visual content lines up."
    )
    _add_snapshot_table(doc, IMG_DIR)

    # ----- 8. Failure analysis -----
    add_heading(doc, "8. Failure analysis", level=1)
    add_para(
        doc,
        "DCT degenerate-luma failure (§7 baseline 3 %): QIM on a flat Y=0 "
        "or Y=255 block is mathematically impossible — the spatial "
        "perturbation needed to write a bit clips at the value floor/ceiling, "
        "and the surviving coefficient returns to ≈δ/2 (the QIM decision "
        "boundary), so any noise flips the bit. The same affects solid "
        "backgrounds, letterbox bars, and blown-out highlights; bounded "
        "as long as enough content-bearing blocks remain for BCH to cover "
        "the lost chips. Fix: a luma-variance gate before embedding."
    )
    add_para(
        doc,
        "Genuine algorithmic limits (§7 attack table): heavy crop (0.6×) "
        "destroys both — DCT loses coordinate alignment; TrustMark's "
        "learned receptive field saturates once 40 % of the frame is "
        "discarded. Aggressive recompression (CRF 36) erases mid-band DCT "
        "coefficients faster than BCH(127, t=11) can recover, and even "
        "TrustMark's residual mostly fails. The screen-capture path (§7.1) "
        "is a separate failure mode entirely — out-of-distribution for "
        "both models."
    )

    # ----- 9. Tradeoff discussion -----
    add_heading(doc, "9. Tradeoff discussion", level=1)
    add_table(doc,
        headers=["Dimension", "DCT-QIM", "TrustMark"],
        rows=[
            ["Visual fidelity (PSNR)", "Higher (~57 dB)", "Lower (~47 dB)"],
            ["Codec robustness (≤ CRF 32)", "Weak",            "Strong"],
            ["Resize / mild-crop robustness", "Brittle",       "Strong"],
            ["Heavy crop / screen-capture", "Brittle",         "Brittle"],
            ["Gaussian noise",          "Strong",              "Strong"],
            ["Compute / latency",        "CPU, ms / frame",    "GPU preferred, ~40 ms / frame"],
            ["Dependencies",             "bchlib only",        "PyTorch + ~200 MB weights"],
            ["Determinism",              "Fully deterministic",  "Stochastic (model precision)"],
        ],
        col_widths=[Inches(2.0), Inches(2.2), Inches(2.2)],
    )
    add_para(
        doc,
        "DCT is the clean choice when invisibility and deployment "
        "simplicity dominate and the only threat is moderate codec "
        "compression or noise. TrustMark wins under any pipeline that may "
        "resize, crop, or transcode the video — paid for with model "
        "weights and inference cost. The two methods' weaknesses are "
        "largely disjoint, so running both in parallel and cross-checking "
        "is the natural high-stakes deployment."
    )

    # ----- 10. Future improvements -----
    add_heading(doc, "10. Future improvements", level=1)
    add_bullets(doc, [
        "Content-aware DCT embedding — skip / scale-down δ on flat-luma "
        "blocks to eliminate the §8 degenerate-luma failure without "
        "dataset curation.",
        "I-frame-aware embedding — PyAV exposes packet flags; raise δ on "
        "I-frames and cut-adjacent frames where the codec is harshest.",
        "Hybrid DCT + TrustMark pipeline — decode both, weighted-vote or "
        "fall-back; weaknesses are disjoint (§9).",
        "Scale- and rotation-equivariant DCT — estimate scale/rotation "
        "via a known marker, rectify before decode. Closes the DCT-TM "
        "gap on resize and recovers the screen-capture case (§7.1).",
        "Train / use a print-scan-tolerant TrustMark variant — addresses "
        "the camera-capture failure mode at the cost of a different "
        "weight file.",
        "Adversarial-attack additions — oracle attack, watermark-aware "
        "denoising, collusion across multiple watermarked copies.",
        "GPU batching for TrustMark — current per-frame Python overhead "
        "dominates; batched inference would turn 'experimental' "
        "throughput into 'production'.",
    ])

    # ----- References -----
    add_heading(doc, "References", level=1)
    add_bullets(doc, [
        "Bui T., Agarwal S., Collomosse J. (2023). \"TrustMark: Universal "
        "Watermarking for Arbitrary Resolution Images.\" arXiv:2311.18297. "
        "https://arxiv.org/abs/2311.18297  —  reference implementation: "
        "https://github.com/adobe/trustmark",
        "Blender Foundation (2012). \"Tears of Steel\" — Mango Open Movie "
        "Project. https://mango.blender.org/  (source clip excerpt used as "
        "the test video throughout this report).",
    ])

    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
